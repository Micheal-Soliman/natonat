from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(r"D:\projects\nat")
SOURCE = ROOT / "Darna_Comprehensive_Website_Audit_SEO_Proposal_EN.docx"


CONFIGS = {
    "current": {
        "output": ROOT / "Darna_Current_Code_Improvements_SEO_Proposal_EN.docx",
        "title": "Darna Website Audit & Current-Code SEO Proposal",
        "subtitle": "Current-Code Improvement, Full SEO Foundation & Commercial Scope",
        "document_title": "Darna Current-Code Improvements and SEO Proposal",
        "commercial_heading": "15. Commercial Proposal - Current-Code Improvement + Full SEO Foundation",
        "remove_headings": {
            "16. Commercial Option B - Custom Code Rebuild + SEO by Design",
            "17. Commercial Comparison and Recommendation",
            "18. Illustrative Break-Even Model",
        },
        "roadmap_remove": "Phase 3B",
        "roadmap_keep": "Phase 3A",
        "executive": (
            "This standalone proposal focuses on correcting measurable leakage in Darna's current platform "
            "and implementing a complete technical and on-page SEO foundation. The programme retains the "
            "existing application while improving language handling, routing, project discovery, lead capture, "
            "measurement, performance, accessibility and search visibility."
        ),
        "migration_assumption": (
            "The scope is based on improving the current platform and its existing content model. Major platform "
            "replacement, customer portals, online payments, live inventory or unlisted integrations are excluded."
        ),
        "warranty": "Defect warranty: 30 days after launch; new requirements are managed as change requests.",
        "pitch": (
            "Suggested narrative  Darna's website already proves the company exists; the next step is to make it "
            "prove why each project is relevant and capture that intent without leakage. The audit found project "
            "misrouting, mixed-language journeys and placeholder sales options on live forms. It also found that "
            "Darna is visible when users already know the brand, but not consistently when they search by need, "
            "location or unit type. This proposal fixes the current code, strengthens conversion journeys and "
            "implements a complete SEO foundation for 360,000 EGP over 10-12 weeks."
        ),
    },
    "custom": {
        "output": ROOT / "Darna_Custom_Code_Rebuild_SEO_Proposal_EN.docx",
        "title": "Darna Website Audit & Custom-Code SEO Proposal",
        "subtitle": "Custom Code Rebuild, SEO by Design & Commercial Scope",
        "document_title": "Darna Custom-Code Rebuild and SEO Proposal",
        "commercial_heading": "15. Commercial Proposal - Custom Code Rebuild + SEO by Design",
        "remove_headings": {
            "15. Commercial Option A - Current-Code Improvement + Full SEO Foundation",
            "17. Commercial Comparison and Recommendation",
            "18. Illustrative Break-Even Model",
        },
        "roadmap_remove": "Phase 3A",
        "roadmap_keep": "Phase 3B",
        "executive": (
            "This standalone proposal covers a complete custom-code rebuild with SEO designed into the platform "
            "architecture from the start. The programme replaces the current application with a server-rendered, "
            "bilingual platform that gives Darna stronger control over content governance, project discovery, lead "
            "measurement, structured data, performance, accessibility, integrations and future scale."
        ),
        "migration_assumption": (
            "The scope includes migration of the agreed page and content inventory plus a redirect map. Major "
            "customer portals, online payments, live inventory or unlisted integrations are excluded."
        ),
        "warranty": "Defect warranty: 60 days after launch; new requirements are managed as change requests.",
        "pitch": (
            "Suggested narrative  Darna's website already proves the company exists; the next step is to build a "
            "digital platform that turns project demand into measurable opportunities at scale. The audit found "
            "project misrouting, mixed-language journeys, weak non-brand visibility and limited content governance. "
            "This proposal replaces the current platform with a custom, server-rendered bilingual build where SEO, "
            "structured data, CRM measurement, performance and content operations are engineered into the "
            "architecture for 920,000 EGP over 18-22 weeks."
        ),
    },
}


def paragraph_text(element):
    return "".join(element.xpath(".//w:t/text()" )).strip()


def is_heading_one(element):
    if element.tag != qn("w:p"):
        return False
    styles = element.xpath("./w:pPr/w:pStyle/@w:val")
    return bool(styles and styles[0] == "Heading1")


def remove_sections(document, headings):
    body = document._element.body
    deleting = False
    for child in list(body):
        if is_heading_one(child):
            deleting = paragraph_text(child) in headings
        if deleting and child.tag != qn("w:sectPr"):
            body.remove(child)


def set_paragraph_text(paragraph, value):
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def replace_exact_paragraph(document, old, new):
    for paragraph in document.paragraphs:
        if paragraph.text == old:
            set_paragraph_text(paragraph, new)
            return True
    return False


def replace_in_tables(document, old, new):
    replacements = 0
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old in paragraph.text:
                        set_paragraph_text(paragraph, paragraph.text.replace(old, new))
                        replacements += 1
    return replacements


def remove_roadmap_row(document, remove_label, keep_label):
    for table in document.tables:
        for row in list(table.rows):
            row_text = " | ".join(cell.text for cell in row.cells)
            if remove_label in row_text:
                table._tbl.remove(row._tr)
                continue
            if keep_label in row_text:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if keep_label in paragraph.text:
                            set_paragraph_text(paragraph, paragraph.text.replace(keep_label, "Phase 3"))


def build(kind, config):
    document = Document(SOURCE)
    remove_sections(document, config["remove_headings"])

    document.core_properties.title = config["document_title"]
    document.core_properties.subject = config["subtitle"]

    replace_exact_paragraph(document, "Darna Comprehensive Website Audit", config["title"])
    replace_exact_paragraph(
        document,
        "SEO Strategy, Competitive Positioning & Two-Track Commercial Proposal",
        config["subtitle"],
    )
    replace_exact_paragraph(
        document,
        "The recommended commercial sequence is to stop measurable leakage first, then decide whether the business should continue improving the current application or rebuild it as a custom, SEO-led platform. Both implementation options in this report include SEO. The distinction is not 'SEO versus no SEO'; it is tactical correction of the current code versus a new architecture designed for long-term control and scale.",
        config["executive"],
    )
    replace_exact_paragraph(
        document,
        "9.3 SEO workstreams included in both commercial options",
        "9.3 SEO workstreams included in this proposal",
    )

    if kind == "current":
        current_architecture_replacements = {
            "10.2 Recommended custom architecture": "10.2 Architecture safeguards for the current platform",
            "Server-rendered React/Next.js or equivalent framework with predictable HTML output and route-level caching.":
                "Reliable server-rendered or pre-rendered HTML output with predictable route-level caching.",
            "Headless CMS with bilingual content models, validation, preview, approval and scheduled publishing.":
                "Bilingual content validation, preview, approval and scheduled publishing within the current administration layer.",
            "Project API/content model supporting locations, categories, unit types, status, payment data, galleries, documents, FAQs and related content.":
                "A governed project model supporting locations, categories, unit types, status, payment data, galleries, documents, FAQs and related content.",
            "CDN and image transformation service, monitored hosting, automated backups and rollback.":
                "CDN and image optimisation, monitored hosting, automated backups and a tested rollback process.",
            "CRM integration layer with retries, de-duplication, consent and attribution fields.":
                "A reliable CRM handoff with retries, de-duplication, consent and attribution fields.",
            "Automated tests for language routes, canonical/hreflang, forms, redirects and critical CTAs.":
                "Automated regression tests for language routes, canonical/hreflang, forms, redirects and critical CTAs.",
        }
        for old, new in current_architecture_replacements.items():
            replace_exact_paragraph(document, old, new)

    if kind == "current":
        replace_exact_paragraph(
            document,
            "15. Commercial Option A - Current-Code Improvement + Full SEO Foundation",
            config["commercial_heading"],
        )
    else:
        replace_exact_paragraph(
            document,
            "16. Commercial Option B - Custom Code Rebuild + SEO by Design",
            config["commercial_heading"],
        )

    renumber = {
        "19. Acceptance Criteria": "16. Acceptance Criteria",
        "20. Commercial Assumptions and Exclusions": "17. Commercial Assumptions and Exclusions",
        "21. KPIs and Governance": "18. KPIs and Governance",
        "22. Client-Facing Pitch": "19. Client-Facing Pitch",
        "23. Sources and Audited URLs": "20. Sources and Audited URLs",
    }
    for old, new in renumber.items():
        replace_exact_paragraph(document, old, new)

    remove_roadmap_row(document, config["roadmap_remove"], config["roadmap_keep"])

    replace_exact_paragraph(
        document,
        "Option B includes migration of the agreed page and content inventory plus a redirect map; major portals, payments or live inventory are excluded.",
        config["migration_assumption"],
    )
    replace_exact_paragraph(
        document,
        "Defect warranty: 30 days for Option A and 60 days for Option B; new requirements are change requests.",
        config["warranty"],
    )
    replace_in_tables(
        document,
        "Suggested narrative  Darna's website already proves the company exists; the next step is to make it prove why each project is relevant and capture that intent without leakage. The audit found project misrouting, mixed language and placeholder sales options on live journeys. It also found that Darna is visible when users already know the brand, but not consistently when they search by need, location or unit type. Option A fixes the current code and adds a complete SEO foundation. Option B rebuilds the platform with SEO, content governance and CRM measurement designed into the architecture. Both have a separate price, timeline and risk profile.",
        config["pitch"],
    )

    # Ensure no stale option labels remain in either standalone proposal.
    stale = []
    for paragraph in document.paragraphs:
        if "Option A" in paragraph.text or "Option B" in paragraph.text or "Two-Track" in paragraph.text:
            stale.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if "Option A" in cell.text or "Option B" in cell.text or "Two-Track" in cell.text:
                    stale.append(cell.text)
    if stale:
        raise RuntimeError(f"Stale cross-option references in {kind}: {stale}")

    document.save(config["output"])
    return config["output"]


if __name__ == "__main__":
    for proposal_kind, proposal_config in CONFIGS.items():
        print(build(proposal_kind, proposal_config))
