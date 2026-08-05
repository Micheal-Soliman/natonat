from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import math

OUT = Path(r"D:\projects\nat\Darna_Comprehensive_Website_Audit_SEO_Proposal_EN.docx")

NAVY = "142B4A"
BLUE = "245A91"
SKY = "DCE8F4"
GOLD = "C29A4A"
INK = "1D2633"
MUTED = "667085"
LIGHT = "F4F7FA"
RED = "B42318"
AMBER = "B54708"
GREEN = "027A48"
WHITE = "FFFFFF"
FONT = "Arial"


def set_run(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), FONT)
    rpr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    mar = tcpr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tcpr.append(mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa):
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = tcpr.find(qn("w:tcW"))
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        tcpr.append(tcw)
    tcw.set(qn("w:w"), str(dxa))
    tcw.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    total = sum(widths)
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(total))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(indent))
    tblind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        trpr = row._tr.get_or_add_trPr()
        trpr.append(OxmlElement("w:cantSplit"))
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    trpr.append(node)


def add_p(doc, text="", size=10.5, bold=False, color=INK, after=6, before=0,
          align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    set_run(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(text), size={1: 16, 2: 13, 3: 11.5}[level], bold=True,
            color={1: NAVY, 2: BLUE, 3: NAVY}[level])
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    set_run(p.add_run(text), size=10.2)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.10
    set_run(p.add_run(text), size=10.2)
    return p


def add_callout(doc, label, text, fill=SKY, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(label + "  "), size=10.5, bold=True, color=accent)
    set_run(p.add_run(text), size=10.5, color=INK)
    add_p(doc, "", after=3)
    return table


def add_table(doc, headers, rows, widths, font_size=9.1):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    repeat_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(header), size=9.2, bold=True, color=WHITE)
    for ridx, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            if ridx % 2 == 1:
                set_shading(cells[i], LIGHT)
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_run(p.add_run(str(value)), size=font_size)
    set_table_geometry(table, widths)
    add_p(doc, "", after=4)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(paragraph.add_run("Page "), size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_hyperlink(paragraph, label, url):
    rid = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    run.append(rpr)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    link.append(run)
    paragraph._p.append(link)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.78)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

# standard_business_brief preset, with a report-style editorial cover.
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in [
    ("Heading 1", 16, NAVY, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11.5, NAVY, 8, 4),
]:
    style = doc.styles[name]
    style.font.name = FONT
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(header.add_run("DARNA | Digital Audit, SEO & Commercial Proposal"), size=8.5, bold=True, color=MUTED)
add_page_number(section.footer.paragraphs[0])

# Cover
add_p(doc, "DIGITAL EXPERIENCE & SEARCH VISIBILITY REVIEW", size=10, bold=True, color=GOLD,
      after=70, before=38, align=WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc, "Darna Comprehensive Website Audit", size=28, bold=True, color=NAVY,
      after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc, "SEO Strategy, Competitive Positioning & Two-Track Commercial Proposal", size=14,
      color=BLUE, after=24, align=WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc, "Palm Hills Developments used as a strategic reference", size=12.5, bold=True,
      color=INK, after=72, align=WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc, "Audit date: 5 August 2026", size=11, bold=True, color=NAVY,
      after=5, align=WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc, "Scope: UX, conversion, content, SEO, technology, performance, accessibility, analytics, CRM, security and delivery economics",
      size=9.7, color=MUTED, after=28, align=WD_ALIGN_PARAGRAPH.CENTER)
add_callout(doc, "Commercial objective", "Turn the public website from a corporate brochure into a measurable acquisition platform that helps prospects discover the right project, qualify themselves and enter a trackable sales workflow.", fill=LIGHT, accent=NAVY)
doc.add_page_break()

add_heading(doc, "1. Executive Summary", 1)
add_callout(doc, "Bottom line", "Darna has a credible operating story, an established portfolio, bilingual content and several sound SEO foundations. However, high-impact defects in language handling, project routing and lead capture create avoidable trust and conversion loss. Its current search footprint is strong for branded intent but weak for broad, non-branded demand.")
add_p(doc, "The recommended commercial sequence is to stop measurable leakage first, then decide whether the business should continue improving the current application or rebuild it as a custom, SEO-led platform. Both implementation options in this report include SEO. The distinction is not 'SEO versus no SEO'; it is tactical correction of the current code versus a new architecture designed for long-term control and scale.")

scores = [
    ("Brand and trust", "68/100", "Credible history and proof points; limited verified customer proof"),
    ("UX and information architecture", "56/100", "Simple navigation, weak project discovery and comparison"),
    ("Lead generation and conversion", "44/100", "Generic forms, placeholder options and weak qualification"),
    ("Content quality", "61/100", "Useful volume; editing, hierarchy and consistency need work"),
    ("On-page SEO", "66/100", "Canonical, hreflang and schema exist; metadata defects remain"),
    ("Technical SEO", "58/100", "Good foundations; route and language-state risks"),
    ("Performance readiness", "60/100", "Good lazy-loading signal; no verified lab/RUM baseline"),
    ("Accessibility", "57/100", "Image alternatives are strong; forms and carousels need work"),
    ("Analytics maturity", "48/100", "GA present; event and CRM attribution not verified"),
    ("Overall indicative score", "58/100", "Capable base, not yet a mature acquisition platform"),
]
add_table(doc, ["Area", "Score", "Interpretation"], scores, [2500, 1300, 5560], font_size=9.1)
add_p(doc, "The scores are advisory, not Lighthouse scores. No access was provided to the administration panel, source repository, server configuration, GA4, Google Search Console, CRM, revenue data or user research.", size=9.2, color=MUTED)

add_heading(doc, "2. Priority Findings That Support the Business Case", 1)
for item in [
    "English internal pages were observed with Arabic navigation, Arabic footer content and Arabic selected as the active language.",
    "The homepage hero promoted Long Island while its 'Project Details' call-to-action linked to Terrace Mall.",
    "The Terrace Mall lead form exposed 'Option 1' and 'Option 2' for payment system and unit type instead of production data.",
    "The homepage title reads 'Darna Real Estate | darna best Leading Real Estate Developer'; English OG locale was set to Arabic and the OG image used a favicon.",
    "The project catalogue has no useful location/type/status filters and relies on long descriptive cards.",
    "GA4 is installed, but conversion event coverage, UTM persistence, CRM delivery and sales feedback were not verifiable.",
    "Darna was not observed in the first-page organic result sets returned for eight sampled non-brand queries, while brand-led queries surfaced Darna-owned pages strongly.",
]:
    add_bullet(doc, item)

add_heading(doc, "3. Audit Scope, Evidence and Limitations", 1)
add_table(doc, ["Verified from the public experience", "Requires internal access"], [
    ("Homepage, projects index, project detail, contact and careers journeys", "CMS/admin panel, repository and API documentation"),
    ("Metadata, canonical, hreflang, Open Graph, JSON-LD and headings", "GA4, Search Console, CRM and sales disposition data"),
    ("Forms, links, language state, mobile behaviour and calls to action", "Server logs, hosting configuration and security testing"),
    ("Search result sampling and external market lists", "Audited revenue, sales volume and market share data"),
    ("Palm Hills public experience as a benchmark", "User interviews, broker interviews and call-centre recordings"),
], [4680, 4680], font_size=9.2)
add_callout(doc, "Ranking caveat", "There is no single official ranking that places every Egyptian developer in order. A defensible 'Darna is number X' claim requires a defined metric such as audited sales, market share, website traffic, share of search or brand awareness. This report therefore separates industry-list presence from sampled search visibility.", fill="FFF6E5", accent=AMBER)

add_heading(doc, "4. Market Position and Search Ranking Review", 1)
add_heading(doc, "4.1 Industry-list position", 2)
add_p(doc, "Darna was not included in the two reviewed top-10 developer lists. Palm Hills was ranked second in both the Official Egyptian Real Estate Platform's 2025 list and Nawy's 2025 top-10 article updated in July 2026. The Official Platform also stated that Egypt had more than 1,900 developers by 2025, which reinforces why a top-10 absence is not equivalent to a precise national rank.")
industry_rows = [
    ("Official Egyptian Real Estate Platform list", "Not listed in top 10", "Palm Hills #2; TMG #1"),
    ("Nawy top-10 article, updated Jul 2026", "Not listed in top 10", "Palm Hills #2; TMG #1"),
    ("Nawy public developer directory snapshot", "Not observed in captured directory list", "Palm Hills: 69 compounds / 1,463 properties"),
    ("Defensible conclusion", "Outside the reviewed published top-10 sets", "No exact national rank can be claimed"),
]
add_table(doc, ["Source / lens", "Darna", "Reference signal"], industry_rows, [3250, 3050, 3060], font_size=8.9)
add_p(doc, "Marketplace directory coverage can depend on commercial listings and inventory relationships, so it should not be treated as an audited market-share table.", size=9.1, color=MUTED)

add_heading(doc, "4.2 Sampled organic search visibility", 2)
serp_rows = [
    ("Darna real estate developer Egypt", "Darna-owned result observed at the top", "Strong branded relevance"),
    ("best real estate developers in Egypt", "Not observed in returned first-page set", "Authority/list inclusion gap"),
    ("top real estate companies in Egypt", "Not observed in returned first-page set", "Authority/list inclusion gap"),
    ("real estate developer Egypt", "Not observed in returned first-page set", "Broad category gap"),
    ("real estate developer El Shorouk", "Not observed in returned first-page set", "Local category opportunity"),
    ("commercial units for sale El Shorouk", "Not observed in returned first-page set", "Transaction-intent opportunity"),
    ("medical clinics for sale Nasr City", "Darna domain not observed; portals dominated", "Listing and landing-page gap"),
    ("chalets for sale Marsa Matrouh", "Darna domain not observed; portals dominated", "Long Island demand gap"),
    ("real estate investment Egypt developer", "Not observed in returned first-page set", "Informational authority gap"),
]
add_table(doc, ["Sample query", "Observed Darna visibility", "Implication"], serp_rows, [3380, 3130, 2850], font_size=8.6)
add_callout(doc, "Current SEO position", "Brand-led visibility with limited non-brand discoverability. Darna can be found when users already know the name or project, but it is not yet consistently intercepting category, location and transaction demand before a user chooses a developer.", fill=SKY, accent=NAVY)
add_p(doc, "SERP observations are a dated sample, not a permanent rank tracker. Results can vary by location, device, language, personalisation and search-engine updates. Search Console and a neutral rank-tracking tool should establish the contractual baseline.", size=9.1, color=MUTED)

add_heading(doc, "4.3 Competitive search set", 2)
add_table(doc, ["Competitive tier", "Examples", "Why they matter"], [
    ("National brand leaders", "TMG, Palm Hills, Emaar Misr, Mountain View, SODIC, Madinet Masr", "High authority, brand demand, large portfolios and strong earned coverage"),
    ("Local / product competitors", "Z Group, Elegant, Sakan and other East Cairo developers", "Compete for El Shorouk and mixed-use intent"),
    ("Search aggregators", "Bayut, Aqarmap, Property Finder, Nawy", "Own transaction SERPs and expose pricing, filters and inventory"),
    ("Attention competitors", "Social media, broker pages, map results and video platforms", "Capture discovery before a corporate site visit"),
], [2300, 3500, 3560], font_size=8.9)

add_heading(doc, "5. What Darna Already Does Well", 1)
for item in [
    "A credible company story dating to 1990, with measurable trust indicators covering projects, built-up area, clients, employees and delivery rate.",
    "Canonical, hreflang, meta descriptions and index/follow directives on the audited homepage.",
    "WebSite, Organization and LocalBusiness structured data on the homepage; BreadcrumbList on the Terrace Mall page.",
    "All 74 observed homepage images carried alternative text, and 71 were lazy-loaded.",
    "WhatsApp, hotline, email, map location, privacy policy, media content and project brochures are available.",
    "GA4 is installed with measurement ID G-3QPXDEG2RB, creating a base for a proper measurement plan.",
    "Darna-owned pages rank strongly for branded intent and Long Island has an indexable, descriptive project page.",
]:
    add_bullet(doc, item)

add_heading(doc, "6. Detailed Issue Register", 1)
issues = [
    ("P0", "Language-state failure", "English routes expose Arabic navigation/footer and Arabic active state", "Trust, crawling, usability"),
    ("P0", "Wrong hero CTA destination", "Long Island promotion links to Terrace Mall", "Lead leakage and attribution"),
    ("P0", "Production form placeholders", "Payment and unit choices show Option 1 / Option 2", "Lead quality and credibility"),
    ("P1", "Homepage URL variants", "/en, /en/ and /en/home are linked", "Duplicate signals and reporting"),
    ("P1", "Weak metadata", "Unnatural title, Arabic OG locale on English, favicon OG image", "CTR and social sharing"),
    ("P1", "Generic lead capture", "Insufficient qualification and no verified attribution payload", "Sales efficiency"),
    ("P1", "Thin project-commerce layer", "No useful filters, comparison, inventory model or detailed unit presentation", "Discovery and conversion"),
    ("P1", "Non-brand search gap", "No sampled first-page visibility across eight broad queries", "Missed acquisition demand"),
    ("P2", "Dynamic heading inconsistency", "Accessibility snapshot exposed H1 while runtime DOM check did not", "Indexing and accessibility risk"),
    ("P2", "Careers dead end", "No jobs found without general CV, alerts or alternative CTA", "Employer brand loss"),
    ("P2", "English editorial quality", "Grammar and keyword-led phrasing reduce premium perception", "Trust and SEO quality"),
    ("P2", "Carousel semantics", "Repeated partner controls appear multiple times in the accessibility tree", "Screen-reader friction"),
]
add_table(doc, ["Priority", "Issue", "Evidence", "Business impact"], issues, [850, 2000, 3880, 2630], font_size=8.25)

add_heading(doc, "7. UX, Information Architecture and Conversion Audit", 1)
add_heading(doc, "7.1 Homepage", 2)
add_p(doc, "The visual hero is attractive but asks the user to process a long headline, long paragraph, a rotating project and two calls to action. The wrong destination defect makes this area commercially high-risk. Each slide should be a validated project object, not a manually disconnected combination of copy and links.")
for item in [
    "Use one concise value proposition, a project/location/status fact line and one primary conversion action.",
    "Add a mobile sticky action bar for Call, WhatsApp and Request Details.",
    "Replace long project paragraphs with scan-friendly cards and filters for location, asset type and delivery status.",
    "Show verifiable delivery proof, project updates, real customer stories and after-sales evidence.",
    "Remove duplicated semantic carousel controls and respect reduced-motion preferences.",
]:
    add_bullet(doc, item)

add_heading(doc, "7.2 Projects catalogue", 2)
for item in [
    "Filter by location, project type, unit type, status and availability signal.",
    "Allow each card to expose location, status, unit categories and one clear CTA.",
    "Create an indexable landing-page strategy for meaningful filters, not every generated filter combination.",
    "Introduce a map/list switch only if the project portfolio and user evidence justify it.",
]:
    add_bullet(doc, item)

add_heading(doc, "7.3 Project-detail template", 2)
add_p(doc, "A project page must answer where, what, who it is for, available unit types, spaces, delivery status, payment structure, proof and next step within seconds. Terrace Mall currently provides useful description and facts but does not form a complete sales journey.")
for item in [
    "Hero: project name, location, category, delivery status, starting area/price when approved and contextual CTA.",
    "Sections: Overview, Units, Amenities, Location, Masterplan, Gallery, Construction Updates, FAQs and Related Projects.",
    "First-party brochure delivery with optional lead gate, instead of relying on Google Drive as the only asset path.",
    "Project-specific form with project_id, unit_type, payment preference, language and UTM fields.",
    "Visible, factual FAQs with appropriate structured data; no invented availability or offers.",
]:
    add_bullet(doc, item)

add_heading(doc, "7.4 Lead capture and CRM workflow", 2)
add_p(doc, "The target is a lead-qualification workflow, not a generic contact form. The minimum model should include name, contactable phone, email when needed, project, unit type, budget/payment preference, purchase timing, source, campaign, language and consent status.")
add_callout(doc, "Acceptance rule", "Every successful submission creates one de-duplicated lead, stores source/medium/campaign/project/language, fires one analytics event, displays a clear confirmation state and reaches the CRM with an owner and contact SLA.", fill=LIGHT, accent=GREEN)

add_heading(doc, "8. Content and Trust Audit", 1)
for item in [
    "Rewrite the English experience through professional editing; phrases such as 'darna best Leading' and 'a mark of trusted' reduce credibility.",
    "Create one messaging hierarchy: promise, proof, project relevance and action.",
    "Use a governed project-content model with required fields, owners and review dates.",
    "Build proof assets: delivery case studies, verified testimonials, handover timelines, construction updates and after-sales stories.",
    "Convert the Media Center into topic clusters connected to projects, locations, asset types and buyer questions.",
    "Synchronise Arabic and English publishing so one language does not become stale or structurally different.",
    "Create a content approval workflow covering commercial claims, pricing, availability and legal review.",
]:
    add_bullet(doc, item)

add_heading(doc, "9. Full SEO Audit and Growth Strategy", 1)
add_heading(doc, "9.1 On-page and technical findings", 2)
for item in [
    "Homepage canonical is https://darna-drc.com/en with en/ar/x-default hreflang.",
    "Homepage robots directive is index, follow and a meta description is present.",
    "Homepage title is unnatural; English Open Graph locale is Arabic and the shared image is a favicon.",
    "Multiple homepage paths are linked and require response-code, redirect and canonical consolidation.",
    "Structured data foundations exist but project coverage is limited primarily to breadcrumbs.",
    "robots.txt and sitemap.xml could not be verified from the browser environment because the requests were client-blocked; they must be checked from hosting and Search Console.",
]:
    add_bullet(doc, item)

add_heading(doc, "9.2 Keyword and landing-page architecture", 2)
keyword_rows = [
    ("Brand", "Darna DRC, Darna Real Estate, Darna projects", "Home, About, project hub"),
    ("Category", "real estate developer Egypt, property developer Cairo", "Company/category landing pages"),
    ("Location", "real estate developer El Shorouk, projects in East Cairo", "Location hubs"),
    ("Commercial", "commercial units for sale El Shorouk, offices in El Shorouk", "Project/unit pages"),
    ("Medical", "clinics for sale Nasr City, medical units Cairo", "Medical Center 3 cluster"),
    ("Coastal", "chalets for sale Marsa Matrouh, beachfront chalets Matrouh", "Long Island cluster"),
    ("Investment", "real estate investment Egypt, ready-to-move commercial units", "Guides and project pages"),
    ("Informational", "payment plans, unit types, delivery status, buying guides", "Media Center clusters"),
]
add_table(doc, ["Intent cluster", "Illustrative targets", "Primary destination"], keyword_rows, [1800, 4500, 3060], font_size=8.7)
add_p(doc, "Final targets require Search Console exports, keyword volume/difficulty data, Arabic variants, commercial priorities and sales-team validation. Keyword volume must not be invented.", size=9.1, color=MUTED)

add_heading(doc, "9.3 SEO workstreams included in both commercial options", 2)
for item in [
    "Keyword research, intent map and URL-to-keyword mapping for English and Arabic.",
    "Title, H1, meta description and social-sharing templates for all page types.",
    "Canonical, redirect, hreflang and indexation policy with reciprocal language links.",
    "XML sitemaps for core pages, projects, articles and images with trustworthy lastmod values.",
    "Structured data for Organization, LocalBusiness, Breadcrumb, Article and visible FAQs; project types selected conservatively.",
    "Internal-linking rules between locations, project types, articles and conversion pages.",
    "Image SEO: descriptive files, dimensions, responsive sources, modern formats and meaningful alt text.",
    "Search Console, Bing Webmaster Tools, rank tracking and an organic-lead dashboard.",
    "Migration redirect mapping and post-launch crawl validation when URLs change.",
]:
    add_bullet(doc, item)
add_callout(doc, "Important", "Custom code does not create SEO automatically, and the current platform can still rank. SEO performance depends on crawlable server-rendered HTML, content quality, page speed, metadata, internal linking, authority, publishing operations and measurement.", fill="FFF6E5", accent=AMBER)

add_heading(doc, "9.4 Twelve-month organic growth roadmap", 2)
add_table(doc, ["Period", "Primary focus", "Expected output"], [
    ("Month 0-1", "Baseline, fixes, crawl control and analytics", "Clean indexation and reliable measurement"),
    ("Months 2-3", "Project/location architecture and content rewrite", "High-intent landing-page coverage"),
    ("Months 4-6", "Topic clusters, proof content and digital PR", "Non-brand impressions and authority growth"),
    ("Months 7-9", "Content refresh, CRO and internal-link optimisation", "Higher lead yield from organic traffic"),
    ("Months 10-12", "Gap analysis, link reclamation and portfolio expansion", "Compounding visibility and lower acquisition dependence"),
], [1700, 3980, 3680], font_size=8.9)

add_heading(doc, "10. Technology Architecture and Platform Finding", 1)
add_p(doc, "Although the original brief described the website as WordPress, the public audit did not expose standard WordPress fingerprints such as wp-content or a WordPress generator tag. The application delivered Angular-style main and polyfill bundles, hydration data and content from api.darna-drc.com. The most likely public architecture is an Angular/SSR application with a separate API or CMS. This must be confirmed from hosting, admin access and the repository before a fixed scope is signed.")
add_heading(doc, "10.1 Current-application priorities", 2)
for item in [
    "Confirm framework, SSR mode, API ownership, deployment process and rollback mechanism.",
    "Remove unnecessary page payload and avoid embedding large translation/content objects above the fold.",
    "Create typed project, language and form models so content and links cannot become disconnected.",
    "Separate development, staging and production; add release smoke tests and monitoring.",
    "Document content ownership, cache invalidation and emergency publishing procedures.",
]:
    add_bullet(doc, item)

add_heading(doc, "10.2 Recommended custom architecture", 2)
for item in [
    "Server-rendered React/Next.js or equivalent framework with predictable HTML output and route-level caching.",
    "Headless CMS with bilingual content models, validation, preview, approval and scheduled publishing.",
    "Project API/content model supporting locations, categories, unit types, status, payment data, galleries, documents, FAQs and related content.",
    "CDN and image transformation service, monitored hosting, automated backups and rollback.",
    "CRM integration layer with retries, de-duplication, consent and attribution fields.",
    "Automated tests for language routes, canonical/hreflang, forms, redirects and critical CTAs.",
]:
    add_bullet(doc, item)

add_heading(doc, "11. Performance, Accessibility, Security and Privacy", 1)
add_heading(doc, "11.1 Performance", 2)
for item in [
    "Establish Lighthouse/WebPageTest baselines on a mid-tier mobile profile and collect CrUX/RUM where available.",
    "Reduce JavaScript and embedded-data payload; code-split non-critical modules and defer non-essential third parties.",
    "Keep lazy loading for non-critical imagery, add responsive sources, CDN caching and explicit dimensions.",
    "Use a single critical hero asset strategy and prevent carousel work from degrading interaction responsiveness.",
    "Post-launch targets: LCP <= 2.5s, INP <= 200ms and CLS <= 0.1 at p75 when sufficient real-user data exists.",
]:
    add_bullet(doc, item)

add_heading(doc, "11.2 Accessibility", 2)
for item in [
    "Programmatic labels, required state, instructions and field-linked error messages for every input.",
    "One stable H1 in initial HTML and a logical H2/H3 hierarchy.",
    "Keyboard operation, visible focus, 200% zoom, contrast and RTL/LTR testing.",
    "Hide cloned carousel slides from assistive technology and provide pause/previous/next controls.",
    "Template audit against WCAG 2.2 AA, followed by manual assistive-technology checks.",
]:
    add_bullet(doc, item)

add_heading(doc, "11.3 Security and privacy", 2)
for item in [
    "Server-side validation, rate limiting, spam protection and safe file handling.",
    "Consent wording, purpose limitation, retention rules and privacy-policy review by qualified counsel.",
    "Least-privilege access, MFA where supported, dependency management and security headers.",
    "Encrypted backups, restore testing, audit logs and incident ownership.",
    "No sensitive lead data in analytics payloads, URLs or client-side logs.",
]:
    add_bullet(doc, item)

add_heading(doc, "12. Analytics, Attribution and Sales Measurement", 1)
events = [
    ("view_project", "Project detail viewed", "project_id, language, source"),
    ("filter_projects", "Catalogue filter used", "filter_name, filter_value"),
    ("click_call", "Hotline clicked", "page_type, project_id"),
    ("click_whatsapp", "WhatsApp clicked", "page_type, project_id, campaign"),
    ("download_brochure", "Brochure requested/downloaded", "project_id, gated_status"),
    ("generate_lead", "Qualified form submitted", "project_id, unit_type, source; no PII"),
    ("form_error", "Submission failure", "form_id, error_type; no PII"),
    ("sales_qualified", "CRM lead accepted", "offline conversion import"),
]
add_table(doc, ["Event", "Trigger", "Key non-PII parameters"], events, [2200, 3500, 3660], font_size=8.8)
add_p(doc, "A weekly funnel should connect sessions -> project views -> contact actions -> submitted leads -> contacted leads -> sales-qualified leads -> reservations. Website optimisation without downstream sales feedback risks improving form volume while lowering lead quality.")

add_heading(doc, "13. Palm Hills Benchmark: Adopt the Logic, Not the Website", 1)
benchmark = [
    ("Portfolio organisation", "Communities by East/West Cairo, North Coast, Commercial and Alexandria", "Location/type/status filters suited to Darna's portfolio"),
    ("Conversion", "Persistent 'Request a Sales Call' and hotline", "Contextual project CTA plus sticky mobile contact"),
    ("Brand proposition", "Short, memorable hero message", "Shorter hero with factual proof"),
    ("Content ecosystem", "Hills Today, Clubs and Investor Relations", "Focused project, location and proof clusters"),
    ("Retention", "Newsletter subscription", "Project updates and opt-in alerts"),
    ("Mobile", "Compact menu and mobile-led hero", "Responsive components and device QA"),
]
add_table(doc, ["Area", "Palm Hills signal", "Recommended Darna adaptation"], benchmark, [1700, 3980, 3680], font_size=8.7)
add_p(doc, "Palm Hills is not a perfect technical benchmark. The audit observed 169 images without alt text in the runtime DOM, 41 scripts and text-empty H1/H2 elements, and the page was comparatively heavy to load. Its strategic strengths should be adopted without copying its accessibility and performance liabilities.", size=9.2, color=MUTED)

add_heading(doc, "14. Delivery Roadmap", 1)
roadmap = [
    ("Phase 0", "2 weeks", "Access verification, analytics/Search Console baseline, crawl, backlog and fixed SOW"),
    ("Phase 1", "Weeks 1-3", "P0 defects, language state, CTAs, production form data, metadata and tracking"),
    ("Phase 2", "Weeks 4-8", "Project catalogue/detail UX, content model, CRM and SEO foundations"),
    ("Phase 3A", "Weeks 9-12", "Current-code hardening, QA, launch and training"),
    ("Phase 3B", "Weeks 9-22", "Custom build, migration, integration, UAT, launch and warranty"),
    ("Growth", "Months 2-12", "Content, authority, CRO, reporting and search iteration"),
]
add_table(doc, ["Stage", "Timing", "Outcome"], roadmap, [1500, 1600, 6260], font_size=8.9)

add_heading(doc, "15. Commercial Option A - Current-Code Improvement + Full SEO Foundation", 1)
add_callout(doc, "Best fit", "Faster commercial improvement with lower migration risk. This is not a cosmetic-edit package: it includes code correction, conversion work, project templates, analytics and a full one-time SEO foundation. Final scope is subject to confirming the actual application/CMS architecture.", fill=SKY, accent=BLUE)
option_a = [
    ("Discovery, access and architecture verification", "25,000 EGP"),
    ("Critical code, language, routing and CTA fixes", "40,000 EGP"),
    ("Responsive UX/UI improvements", "55,000 EGP"),
    ("Project catalogue/detail templates and filters", "60,000 EGP"),
    ("Lead forms, CRM handoff and analytics", "40,000 EGP"),
    ("Full SEO foundation: research, mapping, metadata, hreflang, schema, sitemaps and internal linking", "65,000 EGP"),
    ("Performance, accessibility, security and privacy hardening", "45,000 EGP"),
    ("QA, launch support and team training", "30,000 EGP"),
    ("TOTAL", "360,000 EGP"),
]
add_table(doc, ["Workstream", "Price"], option_a, [6900, 2460], font_size=9.0)
add_p(doc, "Estimated duration: 10-12 weeks after access and content readiness.", bold=True, color=NAVY)
add_p(doc, "Optional ongoing services: technical support and monitoring at 20,000 EGP/month; SEO growth retainer at 35,000 EGP/month for content planning, optimisation, reporting and authority work. Media spend and large-scale content production are excluded.")

add_heading(doc, "16. Commercial Option B - Custom Code Rebuild + SEO by Design", 1)
add_callout(doc, "Best fit", "A three-year digital platform roadmap requiring stronger performance control, scalable bilingual publishing, structured project data, CRM integration and lower dependence on the current application. SEO is built into discovery, architecture, implementation and migration.", fill=SKY, accent=NAVY)
option_b = [
    ("Discovery, business requirements and SEO strategy", "60,000 EGP"),
    ("UX/UI design system for desktop, mobile, English and Arabic", "110,000 EGP"),
    ("Server-rendered custom frontend", "220,000 EGP"),
    ("Headless CMS, bilingual content models and APIs", "150,000 EGP"),
    ("Projects, filters, lead journeys and brochure workflow", "95,000 EGP"),
    ("SEO architecture, structured data, migration and redirect programme", "95,000 EGP"),
    ("CRM, GA4/GTM and reporting integrations", "55,000 EGP"),
    ("Performance, accessibility, security and full QA", "85,000 EGP"),
    ("Deployment, training and post-launch warranty", "50,000 EGP"),
    ("TOTAL", "920,000 EGP"),
]
add_table(doc, ["Workstream", "Price"], option_b, [6900, 2460], font_size=9.0)
add_p(doc, "Estimated duration: 18-22 weeks after discovery and content readiness.", bold=True, color=NAVY)
add_p(doc, "Optional ongoing services: platform support at 28,000 EGP/month; SEO growth retainer at 35,000 EGP/month. Complex live inventory, customer portals, online payments, broker portals or unlisted integrations require separate scope.")

add_heading(doc, "17. Commercial Comparison and Recommendation", 1)
comparison = [
    ("One-time price", "360,000 EGP", "920,000 EGP"),
    ("Delivery", "10-12 weeks", "18-22 weeks"),
    ("SEO included", "Full foundation + implementation", "Strategy, architecture, implementation + migration"),
    ("Migration risk", "Low to medium", "Medium to high"),
    ("Performance control", "Good, limited by current architecture", "High when engineered correctly"),
    ("Content governance", "Depends on current admin/API", "Designed around Darna's team"),
    ("Integration flexibility", "Medium", "High"),
    ("Recommended use", "Stop leakage and prove conversion", "Long-term platform after roadmap approval"),
]
add_table(doc, ["Criterion", "Option A: current code", "Option B: custom code"], comparison, [2300, 3530, 3530], font_size=8.8)
add_callout(doc, "Recommendation", "Start with a paid two-week Verification & Measurement phase priced at 45,000 EGP and credit it against the selected implementation. This validates the true platform, Search Console baseline, conversion tracking, redirect scope and integration constraints before a fixed-price commitment.", fill="FFF6E5", accent=AMBER)

add_heading(doc, "18. Illustrative Break-Even Model", 1)
add_p(doc, "Because sales values and gross contribution were not supplied, the report does not claim a return on investment. The table below shows only how many incremental deals would be needed to recover the one-time implementation price at different hypothetical gross contributions per completed deal.")
break_even = []
for contribution in (150000, 300000, 500000):
    break_even.append((f"{contribution:,.0f} EGP", math.ceil(360000 / contribution), math.ceil(920000 / contribution)))
add_table(doc, ["Hypothetical contribution / deal", "Option A deals", "Option B deals"], break_even, [3800, 2780, 2780], font_size=9.2)
add_p(doc, "Replace these scenarios with finance-approved contribution and lead-to-sale conversion data before presenting an ROI commitment.", size=9.2, color=MUTED)

add_heading(doc, "19. Acceptance Criteria", 1)
for item in [
    "No P0 language, route, CTA or form defects across the agreed device/browser matrix.",
    "Every indexable page has an approved title, H1, canonical, language annotation and social-sharing preview.",
    "Redirect, sitemap, robots and structured-data checks pass in staging and production.",
    "All agreed lead events fire once, contain no PII and reconcile with CRM submissions.",
    "Forms provide accessible labels, validation, success/failure states and spam protection.",
    "Performance targets are measured against the approved baseline and documented test profile.",
    "Content editors can create and preview bilingual projects without developer intervention for standard fields.",
    "A rollback plan, monitoring dashboard, administrator guide and training session are delivered.",
]:
    add_bullet(doc, item)

add_heading(doc, "20. Commercial Assumptions and Exclusions", 1)
for item in [
    "Prices are before applicable taxes and remain valid for 30 days.",
    "Suggested payment schedule: 40% on start, 30% on design/architecture approval, 20% before UAT and 10% at launch.",
    "The client provides timely access, brand assets, approved project data, legal text, analytics accounts and CRM contacts.",
    "Hosting, third-party licences, paid media, photography/video, legal translation and unlisted bulk content entry are excluded.",
    "Option B includes migration of the agreed page and content inventory plus a redirect map; major portals, payments or live inventory are excluded.",
    "Defect warranty: 30 days for Option A and 60 days for Option B; new requirements are change requests.",
    "Search ranking and revenue cannot be guaranteed. Deliverables cover best-practice implementation, measurement and optimisation capability.",
]:
    add_bullet(doc, item)

add_heading(doc, "21. KPIs and Governance", 1)
kpis = [
    ("Search visibility", "Non-brand impressions, top-10 keyword count, share of search", "Monthly"),
    ("Organic acquisition", "Qualified organic leads and organic lead rate", "Weekly / monthly"),
    ("Conversion", "Project view -> contact action -> form -> SQL", "Weekly"),
    ("Lead quality", "% with project/source/contactable phone", ">=95% target"),
    ("Sales SLA", "Time to first contact and contact rate", "Daily / weekly"),
    ("Experience", "Form completion, error rate and mobile conversion", "Weekly"),
    ("Core Web Vitals", "LCP, INP and CLS at p75", "Monthly when RUM is sufficient"),
    ("Quality", "Broken links, wrong-language pages and failed forms", "Zero P0 in production"),
]
add_table(doc, ["Area", "KPI", "Cadence / target"], kpis, [1800, 4600, 2960], font_size=8.8)

add_heading(doc, "22. Client-Facing Pitch", 1)
add_callout(doc, "Suggested narrative", "Darna's website already proves the company exists; the next step is to make it prove why each project is relevant and capture that intent without leakage. The audit found project misrouting, mixed language and placeholder sales options on live journeys. It also found that Darna is visible when users already know the brand, but not consistently when they search by need, location or unit type. Option A fixes the current code and adds a complete SEO foundation. Option B rebuilds the platform with SEO, content governance and CRM measurement designed into the architecture. Both have a separate price, timeline and risk profile.", fill=SKY, accent=NAVY)

add_heading(doc, "23. Sources and Audited URLs", 1)
sources = [
    ("Darna homepage", "https://darna-drc.com/en"),
    ("Darna projects", "https://darna-drc.com/en/projects"),
    ("Darna Terrace Mall", "https://darna-drc.com/en/projects/terrace-mall-darna"),
    ("Darna contact", "https://darna-drc.com/en/contact-us"),
    ("Darna careers", "https://darna-drc.com/en/careers"),
    ("Palm Hills homepage", "https://www.palmhillsdevelopments.com/en-us/home"),
    ("Official Egyptian Real Estate Platform - Best Developers 2025", "https://blogs.realestate.gov.eg/best-real-estate-developers-in-egypt-2025/"),
    ("Nawy - Top 10 Developers 2025, updated July 2026", "https://www.nawy.com/blog/99712-top-real-estate-developers-sales-achievers"),
    ("Nawy developer directory", "https://www.nawy.com/developer"),
    ("Bayut clinics for sale in Nasr City", "https://www.bayut.eg/en/cairo/clinics-for-sale-in-nasr-city/"),
    ("Aqarmap clinics for sale in Nasr City", "https://aqarmap.com.eg/en/for-sale/clinic/cairo/nasr-city/"),
]
for label, url in sources:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    set_run(p.add_run(label + ": "), size=8.8, bold=True)
    add_hyperlink(p, "Open source", url)
add_p(doc, "All findings represent a public snapshot on the audit date. Search results and website states can change; the signed SOW should use a fresh baseline and documented access review.", size=8.8, color=MUTED, after=2)

doc.core_properties.title = "Darna Comprehensive Website Audit, SEO Strategy and Commercial Proposal"
doc.core_properties.subject = "English website audit with competitive ranking review and separate current-code/custom-code pricing"
doc.core_properties.author = "Digital Experience Audit"
doc.core_properties.keywords = "Darna, SEO, website audit, UX, conversion, custom code, Palm Hills, Egypt real estate"
doc.save(OUT)
print(str(OUT))
