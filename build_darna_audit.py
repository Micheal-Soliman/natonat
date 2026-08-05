from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_ROW_HEIGHT_RULE
from pathlib import Path

OUT = Path(r"D:\projects\nat\Darna_Website_Audit_Proposal_AR.docx")

NAVY = "142B4A"
BLUE = "245A91"
SKY = "DCE8F4"
GOLD = "C29A4A"
INK = "1D2633"
MUTED = "667085"
LIGHT = "F4F7FA"
RED = "B42318"
AMBER = "B54708"
GREEN = "027A48"
WHITE = "FFFFFF"
FONT = "Arial"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_rtl(paragraph, rtl=True):
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if rtl and bidi is None:
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        p_pr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_p(doc, text="", size=10.5, bold=False, color=INK, after=6, before=0,
          rtl=True, align=None, italic=False, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    set_rtl(p, rtl)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    set_rtl(p)
    set_run_font(p.add_run(text), size=10.3)
    return p


def add_num(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    set_rtl(p)
    set_run_font(p.add_run(text), size=10.3)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_rtl(p)
    keep_with_next(p)
    r = p.add_run(text)
    sizes = {1: 16, 2: 13, 3: 11.5}
    colors = {1: NAVY, 2: BLUE, 3: NAVY}
    set_run_font(r, size=sizes[level], bold=True, color=colors[level])
    return p


def add_callout(doc, label, text, fill=SKY, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_rtl(p)
    r = p.add_run(label + "  ")
    set_run_font(r, size=10.5, bold=True, color=accent)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    add_p(doc, "", after=4)


def add_table(doc, headers, rows, widths, font_size=9.3):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(h), size=9.3, bold=True, color=WHITE)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if ridx % 2 == 1:
                set_cell_shading(cells[i], LIGHT)
            p = cells[i].paragraphs[0]
            set_rtl(p)
            if i == 0 and len(headers) > 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.add_run(str(value)), size=font_size, color=INK)
    set_table_geometry(table, widths, indent=120)
    add_p(doc, "", after=4)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("صفحة ")
    set_run_font(run, size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.8)
sec.bottom_margin = Inches(0.78)
sec.left_margin = Inches(0.9)
sec.right_margin = Inches(0.9)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

# Styles: standard_business_brief, adapted to Arabic while keeping exact preset rhythm.
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in [
    ("Heading 1", 16, NAVY, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11.5, NAVY, 8, 4),
]:
    s = doc.styles[name]
    s.font.name = FONT
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(color)
    s._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    s._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    s._element.rPr.rFonts.set(qn("w:cs"), FONT)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

# Running header and footer.
header = sec.header
hp = header.paragraphs[0]
set_rtl(hp)
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_run_font(hp.add_run("DARNA | Website Audit & Growth Proposal"), size=8.5, bold=True, color=MUTED)
footer = sec.footer
fp = footer.paragraphs[0]
add_page_number(fp)

# Editorial cover.
add_p(doc, "DIGITAL EXPERIENCE AUDIT", size=10, bold=True, color=GOLD, rtl=False,
      align=WD_ALIGN_PARAGRAPH.CENTER, after=70, before=40)
add_p(doc, "تقرير تدقيق شامل لموقع دارنا", size=28, bold=True, color=NAVY,
      align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
add_p(doc, "Website Audit, Growth Roadmap & Commercial Proposal", size=14, color=BLUE,
      rtl=False, align=WD_ALIGN_PARAGRAPH.CENTER, after=28)
add_p(doc, "مقارنة مرجعية مع Palm Hills Developments", size=13, bold=True, color=INK,
      align=WD_ALIGN_PARAGRAPH.CENTER, after=70)
add_p(doc, "تاريخ الفحص: 5 أغسطس 2026", size=11, bold=True, color=NAVY,
      align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
add_p(doc, "نطاق الفحص: تجربة المستخدم، المحتوى، التحويل، SEO، التقنية، الوصول، الأداء، والثقة", size=10,
      color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=28)
add_callout(doc, "الهدف التجاري", "تحويل الموقع من واجهة تعريفية إلى منصة تسويق ومبيعات تقيس الـ leads، تبرز المشروعات بوضوح، وتدعم نمو البحث العضوي.", fill=LIGHT, accent=NAVY)
doc.add_page_break()

add_heading(doc, "1. الملخص التنفيذي", 1)
add_callout(doc, "الخلاصة", "الموقع يمتلك أساسًا جيدًا: هوية واضحة، محتوى ثنائي اللغة، مشروعات حقيقية، إشارات ثقة، وتهيئة SEO أولية. لكن توجد أخطاء مباشرة في رحلة المستخدم واللغة والـ lead capture تقلل الثقة والتحويل، وبعضها ظاهر على صفحات البيع نفسها.", fill=SKY, accent=BLUE)
add_p(doc, "التوصية ليست هدم الموقع فورًا. المسار التجاري الأنسب يبدأ بمرحلة إصلاح مركزة للنظام الحالي، ثم اتخاذ قرار مبني على القياس بين الاستمرار عليه أو إعادة البناء. إذا كان هدف العميل إطلاق منصة قابلة للتوسع، سريعة، ومتكاملة مع CRM ومحتوى المشروعات، فالإعادة المبرمجة Custom تكون استثمارًا منطقيًا؛ أما SEO وحده فليس سببًا كافيًا للهجرة.")

score_rows = [
    ("الهوية والثقة", "68/100", "جيد مع نقص في proof assets والتجارب الواقعية"),
    ("UX وهيكلة المعلومات", "56/100", "تصفح بسيط لكن اكتشاف المشروعات ضعيف"),
    ("التحويل وتوليد العملاء", "44/100", "نماذج عامة وغير مؤهلة وبعض الحقول placeholder"),
    ("المحتوى", "61/100", "كمية جيدة، لكن الصياغة والاختصار والبنية تحتاج مراجعة"),
    ("On-page SEO", "66/100", "Canonical وhreflang وschema موجودة مع عيوب في title وOG"),
    ("Technical SEO", "58/100", "إشارات جيدة لكن ازدواج المسارات وحالة اللغة تحتاج حسم"),
    ("الأداء - تقدير بنيوي", "60/100", "تحميل صور lazy جيد؛ لا توجد قياسات lab معتمدة ضمن النطاق"),
    ("Accessibility", "57/100", "alt text جيد؛ النماذج والكاروسيل والدلالات تحتاج تقوية"),
    ("النتيجة الإجمالية التقديرية", "58/100", "موقع قابل للتحسين بسرعة لكنه لا يعمل بعد كمنصة مبيعات ناضجة"),
]
add_table(doc, ["المحور", "الدرجة", "القراءة"], score_rows, [2500, 1300, 5560], font_size=9.2)
add_p(doc, "ملاحظة منهجية: الدرجات استشارية وليست نتيجة Lighthouse. لم يتوفر وصول إلى لوحة الإدارة، الخادم، GA4، Google Search Console، CRM أو بيانات المبيعات؛ لذلك تم فصل الحقائق المرصودة عن البنود التي تحتاج تحققًا داخليًا.", size=9.3, color=MUTED)

add_heading(doc, "2. أهم نقاط القرار للعميل", 1)
for t in [
    "إصلاح حالة اللغة على الصفحات الإنجليزية فورًا؛ لأن القوائم والتذييل ظهرت بالعربية في صفحات Projects وContact وCareers وصفحة مشروع.",
    "تصحيح رابط CTA في بطل الصفحة الرئيسية؛ زر Project Details أثناء عرض Long Island وجّه إلى Terrace Mall.",
    "تحويل صفحات المشروعات إلى صفحات بيع حقيقية: تفاصيل الوحدات، المزايا، الموقع، المخطط، أنظمة السداد، حالة التسليم، gallery، brochure، وCTA مرتبط بالمشروع.",
    "استبدال Option 1 / Option 2 في نموذج المشروع ببيانات فعلية، وربط كل lead بمصدر الحملة والمشروع والوحدة واللغة.",
    "تنظيف SEO: title طبيعي، OG locale وصورة مشاركة مناسبة، توحيد /en و/en/ و/en/home، وفحص sitemap/robots من الخادم.",
    "تقوية القياس: أحداث GA4 وUTM وWhatsApp/call/form tracking ولوحة funnel أسبوعية.",
]:
    add_bullet(doc, t)

add_heading(doc, "3. النطاق والمنهجية", 1)
add_p(doc, "تم فحص الصفحات العامة للموقعين باللغتين والمسارات الأساسية، ومراجعة DOM وmetadata وschema والنماذج والروابط وهيكل العناوين والصور وإشارات التحليلات. استُخدم Palm Hills كمرجع للـ content architecture ووضوح sales CTAs وتنظيم المجتمعات، وليس كمعيار تقني مثالي.")
add_table(doc, ["تم التحقق منه", "لم يتوفر ضمن الفحص"], [
    ("Home، Projects، Project detail، Contact، Careers", "لوحة الإدارة وقاعدة البيانات"),
    ("Titles، descriptions، canonical، hreflang، OG، JSON-LD", "GA4 وSearch Console وCRM data"),
    ("هيكل النماذج والروابط واللغة والـ CTAs", "Lighthouse lab / RUM وserver logs"),
    ("مقارنة experience مع Palm Hills", "اختبار مستخدمين ومكالمات المبيعات"),
], [4680, 4680], font_size=9.4)

add_heading(doc, "4. ما يعمل جيدًا في موقع Darna", 1)
for t in [
    "قصة خبرة منذ 1990 وأرقام ثقة واضحة: سنوات العمل، المشروعات، المساحات، العملاء، الموظفون، ومعدل التسليم.",
    "وجود canonical وhreflang للإنجليزية والعربية، وmeta descriptions، وJSON-LD من أنواع WebSite وOrganization وLocalBusiness على الصفحة الرئيسية.",
    "جميع صور الصفحة الرئيسية المرصودة تحمل alt text، و71 من 74 صورة كانت lazy-loaded؛ نقطة إيجابية للأداء والوصول.",
    "وجود WhatsApp وhotline والبريد والموقع الجغرافي وسياسة الخصوصية ومركز مقالات.",
    "صفحة Terrace Mall تحمل H1 واضحًا وBreadcrumb schema وبيانات أساسية ورابط ملف المشروع والموقع.",
    "Google Analytics موجود بالمعرّف G-3QPXDEG2RB، ما يتيح بناء measurement plan بدل البدء من الصفر.",
]:
    add_bullet(doc, t)

add_heading(doc, "5. سجل المشكلات ذات الأولوية", 1)
issue_rows = [
    ("P0", "اختلاط اللغة", "صفحات /en الداخلية تعرض menu/footer عربيًا وArabic selected؛ يضر الثقة والتنقل والفهرسة."),
    ("P0", "CTA خاطئ", "Project Details في hero Long Island يشير إلى Terrace Mall؛ lead يذهب للمشروع الخطأ."),
    ("P0", "نموذج مشروع غير مكتمل", "Payment System وUnit Type ظهرا Option 1 / Option 2 بدل بيانات فعلية."),
    ("P1", "ازدواج مسارات Home", "روابط إلى /en و/en/ و/en/home؛ يلزم redirect/canonical policy موحدة."),
    ("P1", "Metadata غير احترافية", "Title: darna best Leading، وog:locale عربي على English، وصورة OG هي favicon."),
    ("P1", "نماذج ضعيفة الدلالة", "لا required/field names/autocomplete ظاهرة في DOM المرصود؛ يؤثر UX والربط والتحليلات."),
    ("P1", "صفحات المشروعات ليست sales-ready", "غياب filters/comparison، والبطاقات تعتمد فقرات طويلة؛ تفاصيل الوحدات محدودة."),
    ("P1", "قياس التحويل غير واضح", "GA موجود، لكن يلزم إثبات Events للاتصال وWhatsApp وbrochure وforms وproject interest."),
    ("P2", "Hero semantics متغيرة", "Accessibility snapshot رصد H1 بينما DOM runtime لم يرصد H1؛ يلزم ضمان H1 ثابت SSR."),
    ("P2", "Careers dead-end", "No jobs found بلا CTA لرفع CV عام أو job alerts أو LinkedIn fallback."),
    ("P2", "صياغة إنجليزية", "عبارات نحوية وتسويقية ضعيفة تقلل الإحساس premium وتشتت الكلمات المفتاحية."),
    ("P2", "Carousel accessibility", "تكرار أزرار partner carousel عدة مرات في الشجرة الدلالية؛ يلزم aria-hidden للنسخ."),
]
add_table(doc, ["الأولوية", "المشكلة", "الأثر التجاري"], issue_rows, [900, 2100, 6360], font_size=8.8)

add_heading(doc, "6. تدقيق UX وConversion", 1)
add_heading(doc, "6.1 الصفحة الرئيسية", 2)
add_p(doc, "الـ hero جذاب بصريًا لكنه يحمّل أكثر من وظيفة: عنوان طويل، وصف طويل، مشروع متغير، وزران. الخطأ في وجهة Project Details يجعل هذه المنطقة عالية المخاطر. الأفضل أن يصبح كل slide كيانًا مرتبطًا بمشروعه، مع CTA أساسي واحد ثابت مثل Request a Call وCTA ثانوي View Project.")
for t in [
    "اختصار العنوان والوصف فوق الطية، وإظهار المشروع والموقع ونوع الوحدات وحالة التسليم في سطر facts.",
    "إضافة sticky conversion bar على الموبايل: Call / WhatsApp / Request details.",
    "عرض 3-4 مشروعات ببطاقات مختصرة بدل فقرات طويلة، مع فلاتر location / type / status.",
    "إضافة testimonials موثقة، صور تسليم فعلية، construction updates، وشعارات الشركاء دون تكرار دلالي.",
]:
    add_bullet(doc, t)

add_heading(doc, "6.2 صفحات المشروعات", 2)
add_p(doc, "صفحة المشروع يجب أن تجيب خلال أقل من 30 ثانية: أين؟ ماذا أشتري؟ المساحات؟ التسليم؟ السداد؟ لماذا هذا المشروع؟ وكيف أتواصل؟ صفحة Terrace Mall الحالية تقدم وصفًا وfacts أساسية لكنها لا تبني رحلة قرار كاملة.")
for t in [
    "Project hero: الاسم، المكان، نوع المشروع، حالة التسليم، starting space/price إن كان مسموحًا، وCTA.",
    "Tabs أو anchors: Overview، Units، Amenities، Location، Masterplan، Gallery، Updates، FAQs.",
    "نموذج سياقي يحمل project_id وunit_type وcampaign/UTM تلقائيًا، ويعرض consent واضحًا.",
    "Brochure downloadable بعد lead capture اختياري، مع حدث analytics؛ لا تجعل الملف على Google Drive هو المصدر الوحيد.",
    "إضافة FAQs حقيقية وschema مناسب، وحالة المخزون/التسليم دون وعود غير قابلة للتحديث.",
]:
    add_bullet(doc, t)

add_heading(doc, "6.3 النماذج والـ CRM", 2)
add_p(doc, "النموذج الحالي يحتاج الانتقال من Contact Form إلى Lead Qualification. الحد الأدنى: الاسم، الهاتف، البريد اختياري حسب استراتيجية المبيعات، المشروع، نوع الوحدة، الميزانية/نظام السداد، التوقيت المتوقع للشراء، مصدر الحملة، اللغة، والموافقة على التواصل.")
add_callout(doc, "قبول التسليم", "كل إرسال ناجح ينشئ lead واحدًا فقط، يحمل source/medium/campaign/project/language، يظهر له thank-you واضح، ويُرسل إلى CRM مع owner وSLA للاتصال.", fill=LIGHT, accent=GREEN)

add_heading(doc, "7. تدقيق المحتوى والثقة", 1)
add_p(doc, "المحتوى الحالي غني نسبيًا لكنه يحتاج تحريرًا احترافيًا. توجد جمل إنجليزية مثل premium materials a mark of trusted وDarna best Leading؛ هذه الصياغات تقلل الثقة أكثر مما تضيف كلمات مفتاحية. المطلوب content system لا مجرد إعادة كتابة الصفحة الرئيسية.")
for t in [
    "Brand messaging hierarchy: وعد واحد، 3 أسباب تصديق، evidence، ثم CTA.",
    "Project content template موحد لكل مشروع مع حقول إلزامية ومراجعة قانونية للمعلومات المتغيرة.",
    "تحويل Media Center إلى clusters مرتبطة بالمشروعات والمواقع وأنواع الوحدات، مع internal linking واضح.",
    "إضافة case studies: مشروع، التحدي، الإنجاز، أرقام التسليم، صور قبل/بعد، وشهادة عميل قابلة للتحقق.",
    "توحيد tone of voice والترجمة؛ لا تكون العربية والإنجليزية مجرد نسخ غير متزامنة.",
]:
    add_bullet(doc, t)

add_heading(doc, "8. تدقيق SEO", 1)
add_heading(doc, "8.1 ما تم رصده", 2)
for t in [
    "Home canonical = https://darna-drc.com/en مع hreflang en/ar/x-default.",
    "Home meta description موجودة، وrobots meta = index, follow.",
    "Home تحمل WebSite وOrganization وLocalBusiness schema، وTerrace Mall تحمل BreadcrumbList.",
    "Home title الحالي غير طبيعي لغويًا، وOG locale مضبوط ar رغم أن الصفحة إنجليزية، وOG image يشير إلى favicon.",
    "توجد روابط متعددة للصفحة الرئيسية: /en، /en/، /en/home؛ يلزم اختبار الاستجابات وتثبيت مسار واحد.",
    "تعذر التحقق من robots.txt وsitemap.xml من بيئة المتصفح بسبب حجب عميل؛ يجب التحقق منهما مباشرة من الخادم/Search Console قبل الإطلاق.",
]:
    add_bullet(doc, t)

add_heading(doc, "8.2 خطة SEO المقترحة", 2)
for t in [
    "Keyword map حسب intent: brand، developer، project، location، unit type، investment، informational.",
    "Title/H1/meta templates منفصلة للصفحات، ومنع حشو الكلمات أو تكرار اسم Darna دون قيمة.",
    "توحيد URLs و301 redirects وcanonical self-reference وhreflang reciprocal.",
    "XML sitemaps منفصلة للصفحات والمشروعات والمقالات والصور، مع lastmod حقيقي.",
    "Schema: Organization/LocalBusiness/Breadcrumb/Article وFAQ عند وجود FAQ مرئي؛ اختيار نوع المشروع بعناية وعدم اختلاق offers.",
    "Image SEO: أسماء وصفية، width/height، WebP/AVIF، srcset، preload للصورة الحرجة فقط، وalt يصف المعنى لا الكلمات المفتاحية.",
    "ربط Search Console وBing Webmaster، dashboard للـ non-brand clicks والـ leads من organic.",
]:
    add_bullet(doc, t)
add_callout(doc, "تصحيح مهم", "WordPress يمكن أن يحقق SEO ممتازًا، وCustom Code يمكن أن يفشل. العامل الحاسم هو بنية المحتوى، SSR/HTML القابل للفهرسة، الأداء، الـ metadata، internal linking، والقدرة على النشر والقياس.", fill="FFF6E5", accent=AMBER)

add_heading(doc, "9. التقنية، الأداء، الوصول والأمان", 1)
add_heading(doc, "9.1 ملاحظة المنصة", 2)
add_p(doc, "رغم أن brief وصف الموقع كـ WordPress، الفحص الخارجي لم يجد meta generator أو مسارات wp-content. الملفات main-*.js وpolyfills-*.js وبيانات __nghData__ تشير إلى Angular/SSR، مع API ظاهر على api.darna-drc.com. يجب تأكيد المعمارية من الاستضافة قبل توقيع نطاق ثابت؛ قد تكون هناك لوحة CMS منفصلة خلف الـ API.")
add_heading(doc, "9.2 الأداء", 2)
for t in [
    "إجراء Lighthouse وWebPageTest على 4G mid-tier، ثم قياس CWV من CrUX/GA4 إن توفر.",
    "تقليل payload؛ الصفحة تحمل scripts وبيانات ترجمة/محتوى كبيرة داخل الاستجابة. فصل ما لا يلزم فوق الطية.",
    "استمرار lazy loading للصور غير الحرجة، مع responsive images وCDN وcache headers وpreconnect محدود.",
    "إيقاف carousels غير الضرورية أو تجميدها بعد التفاعل، واحترام prefers-reduced-motion.",
    "ميزانية قبول: LCP ≤ 2.5s، INP ≤ 200ms، CLS ≤ 0.1 عند p75 من الزيارات الحقيقية بعد توفر بيانات كافية.",
]:
    add_bullet(doc, t)
add_heading(doc, "9.3 Accessibility", 2)
for t in [
    "ربط labels حقيقية بكل input، وإضافة required/aria-describedby ورسائل خطأ مرتبطة بالحقل.",
    "ضمان H1 واحد ثابت في HTML الأولي، وتسلسل H2/H3 منطقي.",
    "إخفاء نسخ carousel المستنسخة عن قارئ الشاشة، وتوفير pause/previous/next بأسماء واضحة.",
    "اختبار keyboard focus والـ contrast والـ zoom 200% وRTL/LTR؛ ثم audit WCAG 2.2 AA على القوالب الأساسية.",
]:
    add_bullet(doc, t)
add_heading(doc, "9.4 Security & Privacy", 2)
for t in [
    "حماية النماذج من spam/rate abuse، validation server-side، وlogging دون تخزين بيانات زائدة.",
    "عرض purpose/consent واضح قبل إرسال lead، ومراجعة سياسة الخصوصية ومدة الاحتفاظ مع مستشار قانوني.",
    "تفعيل security headers، تحديث dependencies، نسخ احتياطية، صلاحيات أقل امتيازًا، وبيئة staging منفصلة.",
]:
    add_bullet(doc, t)

add_heading(doc, "10. Palm Hills كمرجع: ماذا نقتبس وماذا لا نقتبس", 1)
benchmark_rows = [
    ("تنظيم العرض", "تصنيف Communities حسب East/West Cairo، North Coast، Commercial، Alexandria", "إضافة filters واضحة لمشروعات Darna"),
    ("التحويل", "Request a Sales Call ظاهر في الـ header مع hotline", "CTA ثابت + نموذج مشروع contextual"),
    ("العلامة", "Hero قصير ووعد واحد The Pride of an Address", "اختصار Darna hero وإبراز proof"),
    ("المحتوى", "Hills Today، Clubs، Investor Relations", "بناء content pillars تناسب حجم Darna"),
    ("الاحتفاظ", "Newsletter subscription", "تنبيهات مشروعات/تحديثات باختيار المستخدم"),
    ("الموبايل", "Menu مختصر وتجربة hero مهيأة للشاشة الصغيرة", "Sticky CTA واختبار breakpoints فعلي"),
]
add_table(doc, ["المجال", "Palm Hills", "التطبيق المقترح في Darna"], benchmark_rows, [1600, 3850, 3910], font_size=8.9)
add_p(doc, "لا يجب نسخ Palm Hills حرفيًا. الفحص رصد 169 صورة بلا alt text في DOM المرصود، 41 script، وعناوين H1/H2 فارغة نصيًا في runtime، كما أن التحميل كان ثقيلًا. المرجع قوي تسويقيًا وتنظيميًا، لكنه ليس baseline تقنيًا أو Accessibility benchmark.", size=9.5, color=MUTED)

add_heading(doc, "11. خارطة التنفيذ ذات الأولوية", 1)
add_heading(doc, "أول 14 يومًا - وقف التسريب", 2)
for t in [
    "إصلاح اللغة والـ selected state والروابط العربية داخل /en.",
    "تصحيح hero CTA لكل slide واختبار كل رابط خارجي وداخلي.",
    "استبدال placeholder options وربط النماذج بالمشروع والـ UTM.",
    "تنظيف title/OG/locale وصورة المشاركة وتوحيد home URLs.",
    "إعداد GA4 events الأساسية وthank-you states.",
]:
    add_bullet(doc, t)
add_heading(doc, "الأيام 15-45 - رفع التحويل", 2)
for t in [
    "إعادة تصميم Project listing وProject detail template.",
    "تحرير المحتوى الإنجليزي والعربي وبناء component library خفيف.",
    "إضافة filters، sticky CTAs، lead qualification، وCRM workflow.",
    "Performance/accessibility fixes واختبار device matrix.",
]:
    add_bullet(doc, t)
add_heading(doc, "الأيام 46-90 - النمو", 2)
for t in [
    "Topic clusters وlanding pages مدروسة للمواقع وأنواع الوحدات.",
    "Construction updates، case studies، FAQ، ومحتوى proof.",
    "Dashboard أسبوعي للـ funnel، واختبارات A/B على CTA والنموذج.",
    "قرار معماري نهائي: الاستمرار على النظام الحالي أو Custom rebuild وفق البيانات.",
]:
    add_bullet(doc, t)

add_heading(doc, "12. العرض التجاري - المسار A: تطوير النظام الحالي", 1)
add_callout(doc, "مناسب عندما", "نريد نتائج أسرع، أقل مخاطرة، والمحافظة على المحتوى الحالي. السعر يفترض إمكانية التعديل على الـ CMS/API الحالي؛ يتم تثبيته بعد مراجعة admin/hosting/code access.", fill=SKY, accent=BLUE)
wp_rows = [
    ("Discovery + IA + technical verification", "25,000"),
    ("Bilingual UX/content restructuring", "45,000"),
    ("Responsive UI improvements", "55,000"),
    ("Language, routing and CTA fixes", "30,000"),
    ("Lead forms, CRM handoff and analytics", "35,000"),
    ("Technical SEO, schema and redirects", "45,000"),
    ("Performance, accessibility and security hardening", "35,000"),
    ("QA, training and launch support", "25,000"),
    ("الإجمالي المقترح", "295,000 EGP"),
]
add_table(doc, ["البند", "السعر التقديري"], wp_rows, [7000, 2360], font_size=9.4)
add_p(doc, "المدة: 8-10 أسابيع. دعم اختياري بعد الإطلاق: 18,000 EGP شهريًا أو 180,000 EGP سنويًا، يشمل صيانة، مراقبة، إصلاحات صغيرة، وتقرير أداء؛ لا يشمل حملات إعلانية أو إنتاج محتوى كبير.", bold=True, color=NAVY)

add_heading(doc, "13. العرض التجاري - المسار B: إعادة بناء Custom", 1)
add_p(doc, "معمارية مقترحة: Next.js أو إطار SSR حديث + Headless CMS + API/CRM integrations + CDN. الاختيار النهائي بعد discovery. الهدف: HTML قابل للفهرسة، سرعة وتحكم أعلى، قوالب محتوى منظمة، ونشر آمن ثنائي اللغة.")
custom_rows = [
    ("Discovery, requirements and information architecture", "45,000"),
    ("UX/UI design system - desktop/mobile/RTL", "90,000"),
    ("SSR frontend development", "185,000"),
    ("Headless CMS, content models and APIs", "135,000"),
    ("Projects, filters, lead journeys and brochure flow", "85,000"),
    ("SEO migration, redirects, schema and sitemaps", "55,000"),
    ("CRM, GA4/GTM and dashboard instrumentation", "40,000"),
    ("QA, security, deployment, training and warranty", "65,000"),
    ("التكلفة الأساسية المحسوبة", "700,000 EGP"),
    ("النطاق التجاري المتوقع", "700,000-950,000 EGP"),
]
add_table(doc, ["البند", "السعر التقديري"], custom_rows, [7000, 2360], font_size=9.2)
add_p(doc, "عرض تسعير عملي للعميل: 790,000 EGP للنطاق القياسي الموضح، بمدة 16-20 أسبوعًا. يتحول السعر إلى الحد الأعلى عند إضافة integrations معقدة، bulk migration كبير، portal للعملاء، أو وظائف availability/pricing لحظية.", bold=True, color=NAVY)

add_heading(doc, "14. مقارنة القرار", 1)
decision_rows = [
    ("وقت الوصول للسوق", "8-10 أسابيع", "16-20 أسبوعًا"),
    ("التكلفة الأولية", "295K EGP", "790K EGP مقترح"),
    ("مخاطر الهجرة", "منخفضة-متوسطة", "متوسطة-مرتفعة"),
    ("مرونة الأداء", "جيدة حسب المعمارية الحالية", "عالية عند التنفيذ الصحيح"),
    ("سهولة إدارة المحتوى", "تعتمد على لوحة النظام الحالية", "تصمم خصيصًا لفريق Darna"),
    ("SEO", "يمكن أن يكون ممتازًا", "يمكن أن يكون ممتازًا؛ ليس تلقائيًا"),
    ("التوسع والـ integrations", "متوسط", "عالٍ"),
    ("التوصية", "ابدأ به لوقف التسريب", "نفذه إذا وُجدت roadmap 3 سنوات وتمويل وتشغيل واضح"),
]
add_table(doc, ["المعيار", "تطوير الحالي", "Custom rebuild"], decision_rows, [2300, 3530, 3530], font_size=9.0)
add_callout(doc, "التوصية التجارية", "ابدأ بمرحلة Verification & Quick Wins مدفوعة لمدة أسبوعين بقيمة 35,000 EGP وتُخصم من المسار المختار. بعدها يُثبت السعر والنطاق على أساس الوصول الفعلي والقياسات.", fill="FFF6E5", accent=AMBER)

add_heading(doc, "15. الشروط والافتراضات", 1)
for t in [
    "الأسعار قبل أي ضرائب واجبة التطبيق، وصالحة لمدة 30 يومًا.",
    "الدفع المقترح: 40% بدء، 30% اعتماد التصميم، 20% قبل UAT، 10% عند الإطلاق.",
    "يلتزم العميل بتوفير الوصول، المحتوى المعتمد، بيانات المشروعات، وحسابات analytics/CRM في الوقت المتفق عليه.",
    "لا تشمل الأسعار: hosting ورسوم المنصات والتراخيص المدفوعة، تصوير/فيديو، ترجمة قانونية، إدخال محتوى غير متفق عليه، أو حملات media buying.",
    "تتضمن إعادة البناء migration للصفحات المتفق عليها وredirect map؛ أي portal أو payment أو live inventory يسعّر منفصلًا.",
    "ضمان إصلاح العيوب: 30 يومًا بعد الإطلاق للنظام الحالي، و60 يومًا للمسار Custom، ولا يشمل تغييرات scope.",
]:
    add_bullet(doc, t)

add_heading(doc, "16. KPIs التي تُعرض على العميل", 1)
kpi_rows = [
    ("Conversion", "Qualified leads / sessions", "Baseline أول 30 يومًا ثم +25% خلال 90 يومًا"),
    ("Lead quality", "% leads تحمل project + source + contactable phone", "≥ 95%"),
    ("Sales SLA", "متوسط وقت أول اتصال", "يتفق مع فريق المبيعات؛ يُقاس أسبوعيًا"),
    ("Organic", "Non-brand clicks and qualified organic leads", "نمو ربع سنوي بعد تثبيت baseline"),
    ("UX", "Form completion rate / abandonment", "تحسن مستمر عبر الاختبارات"),
    ("CWV", "LCP / INP / CLS p75", "≤2.5s / ≤200ms / ≤0.1 عند توفر RUM"),
    ("Quality", "Broken links, wrong-language pages, failed forms", "صفر P0 في production"),
]
add_table(doc, ["المجال", "المؤشر", "هدف العمل"], kpi_rows, [1700, 3700, 3960], font_size=8.9)

add_heading(doc, "17. نص مختصر لإقناع العميل", 1)
add_callout(doc, "Pitch", "الموقع الحالي يعرّف الناس بدارنا، لكنه لا يستخرج القيمة الكاملة من كل زيارة. المشكلة ليست شكلًا فقط: وجدنا انتقالًا للمشروع الخطأ، لغة عربية داخل صفحات إنجليزية، وخيارات وهمية داخل نموذج البيع. سنوقف هذه التسريبات أولًا، ثم نبني رحلة مشروع قابلة للقياس. Palm Hills يوضح كيف تتحول الهوية إلى مجتمعات وSales CTA واضح؛ سنأخذ المنطق المناسب لحجم دارنا، لا نسخة منه. أمامكم مساران واضحان بالتكلفة والعائد والمخاطر.", fill=SKY, accent=NAVY)

add_heading(doc, "18. الأدلة والروابط التي تم فحصها", 1)
sources = [
    ("Darna - Home", "https://darna-drc.com/en"),
    ("Darna - Projects", "https://darna-drc.com/en/projects"),
    ("Darna - Terrace Mall", "https://darna-drc.com/en/projects/terrace-mall-darna"),
    ("Darna - Contact", "https://darna-drc.com/en/contact-us"),
    ("Darna - Careers", "https://darna-drc.com/en/careers"),
    ("Palm Hills - Home", "https://www.palmhillsdevelopments.com/en-us/home"),
    ("Palm Hills - Community example", "https://www.palmhillsdevelopments.com/en-us/residential-properties/palm-hills-alexandria"),
]
for label, url in sources:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    set_rtl(p)
    set_run_font(p.add_run(label + ": "), size=9.6, bold=True, color=INK)
    add_hyperlink(p, url, url)
add_p(doc, "جميع الملاحظات تمثل snapshot علنيًا بتاريخ الفحص. قد تتغير الواجهات أو البيانات لاحقًا، ويجب إعادة smoke test قبل توقيع الـ SOW النهائي.", size=9.2, color=MUTED)

# Core properties.
doc.core_properties.title = "Darna Website Audit & Growth Proposal"
doc.core_properties.subject = "Comprehensive website audit with Palm Hills benchmark and two commercial implementation paths"
doc.core_properties.author = "Digital Experience Audit"
doc.core_properties.keywords = "Darna, website audit, SEO, UX, conversion, Palm Hills, proposal"

doc.save(OUT)
print(str(OUT))
